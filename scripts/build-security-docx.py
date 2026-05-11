"""
Build a Word document from a markdown source file.

Renders markdown as a polished Word document with:
  - RTL paragraph direction (Arabic) or LTR (English)
  - Suitable font (Calibri for body, Consolas for code)
  - Heading hierarchy preserved
  - Tables rendered as native Word tables
  - Code blocks rendered in a monospaced, framed style

Usage:
  python scripts/build-security-docx.py                 # Arabic, default paths
  python scripts/build-security-docx.py --english       # English, default paths
  python scripts/build-security-docx.py --src X --out Y
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent

# RTL flag is set in main() based on CLI args; renderers read this module-level
# value to decide paragraph/table direction.
RTL = True


# ─── helpers ────────────────────────────────────────────────────────────────

def set_direction(paragraph) -> None:
    """Mark a paragraph LTR or RTL according to the module-level RTL flag."""
    p_pr = paragraph._p.get_or_add_pPr()
    if RTL:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# Back-compat alias used throughout the renderer.
set_rtl = set_direction


def set_font(run, *, name: str = "Calibri", size: int = 11,
             bold: bool = False, color: RGBColor | None = None,
             mono: bool = False) -> None:
    if mono:
        name = "Consolas"
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Ensure complex-script and east-asian fonts also use the same family
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    colors = {
        1: RGBColor(0x0A, 0x25, 0x40),  # ASH navy
        2: RGBColor(0x0C, 0x3D, 0x6E),
        3: RGBColor(0x1F, 0x49, 0x7D),
        4: RGBColor(0x3B, 0x3B, 0x3B),
    }
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run(text)
    set_font(r, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True


# ─── inline parsing ─────────────────────────────────────────────────────────

INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*)"          # bold
    r"|(`([^`]+)`)"               # inline code
    r"|(\[([^\]]+)\]\(([^)]+)\))",  # link
)


def add_inline_runs(paragraph, text: str) -> None:
    """Parse a single paragraph's worth of inline markdown (bold, code, links)."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_font(r)
        if m.group(2):  # **bold**
            r = paragraph.add_run(m.group(2))
            set_font(r, bold=True)
        elif m.group(4):  # `code`
            r = paragraph.add_run(m.group(4))
            set_font(r, mono=True, size=10, color=RGBColor(0xB0, 0x33, 0x00))
        elif m.group(6):  # [text](url) → render as text + url in parens
            r = paragraph.add_run(m.group(6))
            set_font(r, color=RGBColor(0x0C, 0x3D, 0x6E))
            r.underline = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_font(r)


# ─── block parsing ──────────────────────────────────────────────────────────

def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text)


def add_list_item(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    if RTL:
        p.paragraph_format.right_indent = Cm(0.5 + level * 0.75)
    else:
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.75)
    p.paragraph_format.space_after = Pt(3)
    r0 = p.add_run("• ")
    set_font(r0, bold=True, color=RGBColor(0x0A, 0x25, 0x40))
    add_inline_runs(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Force LTR for code
    p_pr = p._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is not None:
        p_pr.remove(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Light gray background via paragraph border
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "D0D0D0")
        pBdr.append(b)
    p_pr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, mono=True, size=9.5)


def add_table_block(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT if RTL else WD_ALIGN_PARAGRAPH.LEFT
    if RTL:
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            bidiVisual = OxmlElement("w:bidiVisual")
            bidiVisual.set(qn("w:val"), "1")
            tblPr.append(bidiVisual)
    # Header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        set_rtl(p)
        r = p.add_run(h)
        set_font(r, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10.5)
        # Fill header dark navy
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0A2540")
        tcPr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_rtl(p)
            add_inline_runs(p, val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "8")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "0A2540")
    pBdr.append(b)
    p_pr.append(pBdr)


# ─── parser ─────────────────────────────────────────────────────────────────

def parse_markdown(md_text: str) -> list[dict]:
    """Return a flat sequence of typed blocks the renderer understands."""
    blocks: list[dict] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Horizontal rule
        if stripped.strip() in ("---", "***"):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.+?)\s*$", stripped)
        if m:
            blocks.append({"type": "h", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append({"type": "code", "lines": code_lines})
            continue

        # Table
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in stripped.strip().strip("|").split("|")]
            i += 2  # past separator
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # Pad / truncate to header width
                if len(row) < len(header):
                    row += [""] * (len(header) - len(row))
                elif len(row) > len(header):
                    row = row[: len(header)]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.+?)\s*$", stripped)
        if m:
            indent = len(m.group(1)) // 2
            blocks.append({"type": "li", "level": indent, "text": m.group(2)})
            i += 1
            continue

        # Blank line
        if not stripped.strip():
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Plain paragraph (collect consecutive non-empty lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,4}\s|```|---|\s*[-*]\s|\s*\|)", lines[i]
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        blocks.append({"type": "p", "text": " ".join(para_lines)})

    return blocks


# ─── render ─────────────────────────────────────────────────────────────────

def render(blocks: list[dict]) -> Document:
    doc = Document()

    # Default font for the body
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page setup — A4 with reasonable margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    if RTL:
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)

    for block in blocks:
        t = block["type"]
        if t == "h":
            add_heading(doc, block["text"], block["level"])
        elif t == "p":
            add_paragraph(doc, block["text"])
        elif t == "li":
            add_list_item(doc, block["text"], block.get("level", 0))
        elif t == "code":
            add_code_block(doc, block["lines"])
        elif t == "table":
            add_table_block(doc, block["header"], block["rows"])
        elif t == "hr":
            add_hr(doc)
        elif t == "blank":
            pass  # paragraphs already have space_after

    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--english", action="store_true", help="Render LTR English layout")
    parser.add_argument("--src", type=Path, help="Source markdown file")
    parser.add_argument("--out", type=Path, help="Output docx file")
    args = parser.parse_args()

    global RTL
    RTL = not args.english

    if args.src is None:
        args.src = ROOT / "docs" / ("security-overview-en.md" if args.english else "security-overview.md")
    if args.out is None:
        args.out = ROOT / "docs" / ("security-overview-en.docx" if args.english else "security-overview.docx")

    md = args.src.read_text(encoding="utf-8")
    blocks = parse_markdown(md)
    doc = render(blocks)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(f"Wrote {args.out.relative_to(ROOT)}  ({args.out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
